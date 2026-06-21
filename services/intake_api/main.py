from __future__ import annotations

import hashlib
import hmac
import io
import json
import os
import re
import uuid
from pathlib import Path
from typing import Annotated, Any

import boto3
import httpx
from bs4 import BeautifulSoup
from docx import Document
from fastapi import Depends, FastAPI, File, Form, Header, HTTPException, Request, UploadFile
from fastapi.responses import FileResponse
from pydantic import BaseModel, ConfigDict, Field, ValidationError, field_validator
from pypdf import PdfReader

APP_VERSION = "1.0.0"
MAX_FILE_BYTES = int(os.getenv("AMAS_MAX_FILE_BYTES", str(25 * 1024 * 1024)))
ALLOWED_MIME = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "text/plain",
    "text/markdown",
    "text/html",
    "application/rtf",
    "text/rtf",
}
N8N_BASE_URL = os.getenv("N8N_INTERNAL_BASE_URL", "http://n8n:5678").rstrip("/")
N8N_INTAKE_PATH = os.getenv("N8N_INTAKE_PATH", "/webhook/amas/v1/cases")
N8N_REPORT_PATH = os.getenv("N8N_REPORT_PATH", "/webhook/amas/v1/cases/{case_id}/report")
N8N_REVIEW_PATH = os.getenv("N8N_REVIEW_PATH", "/webhook/amas/v1/reviews")
AMAS_API_TOKEN = os.getenv("AMAS_API_TOKEN", "")
AMAS_INTERNAL_TOKEN = os.getenv("AMAS_INTERNAL_TOKEN", "")
S3_ENDPOINT = os.getenv("S3_ENDPOINT", "http://garage:3900")
S3_REGION = os.getenv("S3_REGION", "garage")
S3_BUCKET = os.getenv("S3_BUCKET", "amas")
S3_ACCESS_KEY = os.getenv("S3_ACCESS_KEY", "")
S3_SECRET_KEY = os.getenv("S3_SECRET_KEY", "")

app = FastAPI(title="AMAS Intake and Review API", version=APP_VERSION)


@app.middleware("http")
async def add_security_headers(request: Request, call_next):
    response = await call_next(request)
    response.headers["X-Content-Type-Options"] = "nosniff"
    response.headers["X-Frame-Options"] = "DENY"
    response.headers["Referrer-Policy"] = "no-referrer"
    response.headers["Permissions-Policy"] = "camera=(), microphone=(), geolocation=()"
    if request.url.path.startswith("/v1/") or request.url.path.startswith("/internal/"):
        response.headers["Cache-Control"] = "no-store"
    return response


STATIC_DIR = Path(__file__).resolve().parent / "static"


class LearningOutcome(BaseModel):
    model_config = ConfigDict(extra="forbid")
    outcome_code: str
    description: str
    level_descriptor: str | None = None
    metadata: dict[str, Any] = Field(default_factory=dict)


class Artifact(BaseModel):
    model_config = ConfigDict(extra="forbid")
    artifact_type: str
    original_filename: str | None = None
    mime_type: str | None = None
    size_bytes: int | None = None
    storage_bucket: str | None = None
    storage_key: str | None = None
    sha256: str | None = None
    extracted_text: str
    text_sha256: str | None = None
    extraction_status: str = "provided"
    metadata: dict[str, Any] = Field(default_factory=dict)


class CasePayload(BaseModel):
    model_config = ConfigDict(extra="forbid")
    case_id: uuid.UUID | None = None
    external_ref: str | None = None
    tenant_id: str = "default"
    idempotency_key: str | None = None
    title: str
    module_code: str
    module_title: str | None = None
    study_level: str | int | float | None = None
    credit_value: float | None = None
    assessment_weight: float | None = None
    assessment_type: str
    individual_or_group: str | None = None
    formative_or_summative: str | None = None
    issue_date: str | None = None
    submission_date: str | None = None
    submitter_id: str | None = None
    moderator_id: str | None = None
    declared_ai_mode: str | None = "unspecified"
    evaluation_mode: bool = False
    metadata: dict[str, Any] = Field(default_factory=dict)
    learning_outcomes: list[LearningOutcome]
    artifacts: list[Artifact]

    @field_validator("tenant_id")
    @classmethod
    def validate_tenant_id(cls, value: str) -> str:
        value = value.strip()
        if not re.fullmatch(r"[A-Za-z0-9][A-Za-z0-9._-]{0,79}", value):
            raise ValueError("tenant_id must be a 1-80 character slug")
        return value


class ReviewPayload(BaseModel):
    model_config = ConfigDict(extra="forbid")
    case_id: uuid.UUID
    report_id: uuid.UUID
    finding_id: uuid.UUID | None = None
    reviewer_id: str
    decision: str
    comment: str | None = None
    changes: dict[str, Any] = Field(default_factory=dict)


class GeneratedReportPayload(BaseModel):
    model_config = ConfigDict(extra="forbid")
    case_id: uuid.UUID
    report_id: uuid.UUID
    report_version: int
    tenant_id: str = "default"
    report: dict[str, Any]


def authenticate(x_amas_api_token: Annotated[str | None, Header()] = None) -> None:
    if not AMAS_API_TOKEN:
        raise HTTPException(503, "AMAS_API_TOKEN is not configured")
    if not x_amas_api_token or not hmac.compare_digest(x_amas_api_token, AMAS_API_TOKEN):
        raise HTTPException(401, "Invalid API token")


def authenticate_internal(x_amas_internal_token: Annotated[str | None, Header()] = None) -> None:
    if not AMAS_INTERNAL_TOKEN:
        raise HTTPException(503, "AMAS_INTERNAL_TOKEN is not configured")
    if not x_amas_internal_token or not hmac.compare_digest(x_amas_internal_token, AMAS_INTERNAL_TOKEN):
        raise HTTPException(401, "Invalid internal token")


def s3_client():
    return boto3.client(
        "s3",
        endpoint_url=S3_ENDPOINT,
        region_name=S3_REGION,
        aws_access_key_id=S3_ACCESS_KEY,
        aws_secret_access_key=S3_SECRET_KEY,
    )


def safe_filename(value: str | None) -> str:
    value = Path(value or "document").name
    value = re.sub(r"[^A-Za-z0-9._-]+", "_", value).strip("._")
    return value[:180] or "document"


def extract_text(data: bytes, mime: str, filename: str) -> tuple[str, str]:
    try:
        if mime == "application/pdf" or filename.lower().endswith(".pdf"):
            reader = PdfReader(io.BytesIO(data))
            text = "\n\n".join((page.extract_text() or "") for page in reader.pages)
            status = "extracted" if text.strip() else "ocr_required"
            return text, status
        if mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" or filename.lower().endswith(".docx"):
            doc = Document(io.BytesIO(data))
            parts = [p.text for p in doc.paragraphs]
            for table in doc.tables:
                for row in table.rows:
                    parts.append(" | ".join(cell.text for cell in row.cells))
            return "\n".join(parts), "extracted"
        decoded = data.decode("utf-8", errors="replace")
        if mime == "text/html" or filename.lower().endswith((".html", ".htm")):
            return BeautifulSoup(decoded, "html.parser").get_text("\n"), "extracted"
        if mime in {"application/rtf", "text/rtf"} or filename.lower().endswith(".rtf"):
            decoded = re.sub(r"\\'[0-9a-fA-F]{2}", "", decoded)
            decoded = re.sub(r"\\[a-zA-Z]+-?\d* ?", "", decoded)
            decoded = decoded.replace("{", "").replace("}", "")
        return decoded, "extracted"
    except Exception as exc:  # preserve original object; flag extraction error
        return "", f"extraction_error:{type(exc).__name__}"


async def send_to_n8n(payload: dict[str, Any], token: str) -> dict[str, Any]:
    timeout = httpx.Timeout(connect=10, read=300, write=60, pool=10)
    async with httpx.AsyncClient(timeout=timeout) as client:
        response = await client.post(
            f"{N8N_BASE_URL}{N8N_INTAKE_PATH}",
            headers={"X-AMAS-API-Token": token, "Content-Type": "application/json"},
            json=payload,
        )
    if response.status_code >= 400:
        raise HTTPException(502, detail={"message": "n8n intake failed", "status": response.status_code, "body": response.text[:2000]})
    try:
        return response.json()
    except ValueError as exc:
        raise HTTPException(502, "n8n returned a non-JSON response") from exc


@app.get("/health")
async def health() -> dict[str, Any]:
    return {"status": "ok", "service": "amas-intake-api", "version": APP_VERSION}


@app.get("/")
async def index() -> FileResponse:
    return FileResponse(STATIC_DIR / "index.html")


@app.post("/v1/cases/json", dependencies=[Depends(authenticate)])
async def submit_case_json(payload: CasePayload, x_amas_api_token: Annotated[str, Header()]) -> dict[str, Any]:
    body = payload.model_dump(mode="json")
    body["case_id"] = body.get("case_id") or str(uuid.uuid4())
    return await send_to_n8n(body, x_amas_api_token)


@app.post("/v1/cases", dependencies=[Depends(authenticate)])
async def submit_case_files(
    metadata: Annotated[str, Form()],
    assessment_brief: Annotated[UploadFile, File()],
    rubric: Annotated[UploadFile, File()],
    module_descriptor: Annotated[UploadFile | None, File()] = None,
    supporting_documents: Annotated[list[UploadFile] | None, File()] = None,
    x_amas_api_token: Annotated[str, Header()] = "",
) -> dict[str, Any]:
    try:
        meta = json.loads(metadata)
    except json.JSONDecodeError as exc:
        raise HTTPException(422, f"metadata is not valid JSON: {exc}") from exc

    try:
        case_id = str(uuid.UUID(str(meta["case_id"]))) if meta.get("case_id") else str(uuid.uuid4())
    except (ValueError, TypeError, AttributeError) as exc:
        raise HTTPException(422, "case_id must be a UUID") from exc
    tenant_id = str(meta.get("tenant_id", "default")).strip()
    if not re.fullmatch(r"[A-Za-z0-9][A-Za-z0-9._-]{0,79}", tenant_id):
        raise HTTPException(422, "tenant_id must be a 1-80 character slug")
    meta["case_id"] = case_id
    meta["tenant_id"] = tenant_id
    files: list[tuple[str, UploadFile]] = [
        ("assessment_brief", assessment_brief),
        ("rubric", rubric),
    ]
    if module_descriptor:
        files.append(("module_descriptor", module_descriptor))
    for item in supporting_documents or []:
        files.append(("supporting_document", item))

    artifacts: list[dict[str, Any]] = []
    uploaded_keys: list[str] = []
    client = s3_client()
    try:
        for artifact_type, upload in files:
            data = await upload.read(MAX_FILE_BYTES + 1)
            if len(data) > MAX_FILE_BYTES:
                raise HTTPException(413, f"{upload.filename} exceeds {MAX_FILE_BYTES} bytes")
            mime = upload.content_type or "application/octet-stream"
            if mime not in ALLOWED_MIME:
                raise HTTPException(415, f"Unsupported media type for {upload.filename}: {mime}")
            filename = safe_filename(upload.filename)
            digest = hashlib.sha256(data).hexdigest()
            text, extraction_status = extract_text(data, mime, filename)
            text_digest = hashlib.sha256(text.encode("utf-8")).hexdigest()
            object_key = f"originals/{tenant_id}/{case_id}/{artifact_type}/{digest}-{filename}"
            client.put_object(
                Bucket=S3_BUCKET,
                Key=object_key,
                Body=data,
                ContentType=mime,
                Metadata={"case-id": case_id, "artifact-type": artifact_type, "sha256": digest},
            )
            uploaded_keys.append(object_key)
            artifacts.append({
                "artifact_type": artifact_type,
                "original_filename": filename,
                "mime_type": mime,
                "size_bytes": len(data),
                "storage_bucket": S3_BUCKET,
                "storage_key": object_key,
                "sha256": digest,
                "extracted_text": text,
                "text_sha256": text_digest,
                "extraction_status": extraction_status,
                "metadata": {"source": "multipart_upload"},
            })

        meta["artifacts"] = artifacts
        try:
            validated = CasePayload.model_validate(meta).model_dump(mode="json")
        except ValidationError as exc:
            raise HTTPException(422, detail=exc.errors(include_url=False)) from exc
    except Exception:
        # Invalid submissions should not leave unreferenced originals behind.
        for object_key in uploaded_keys:
            try:
                client.delete_object(Bucket=S3_BUCKET, Key=object_key)
            except Exception:
                pass
        raise

    return await send_to_n8n(validated, x_amas_api_token)


@app.get("/v1/cases/{case_id}/report", dependencies=[Depends(authenticate)])
async def get_report(case_id: str, x_amas_api_token: Annotated[str, Header()]) -> Any:
    try:
        canonical_case_id = str(uuid.UUID(case_id))
    except ValueError as exc:
        raise HTTPException(422, "case_id must be a UUID") from exc
    async with httpx.AsyncClient(timeout=60) as client:
        response = await client.get(
            f"{N8N_BASE_URL}{N8N_REPORT_PATH.format(case_id=canonical_case_id)}",
            headers={"X-AMAS-API-Token": x_amas_api_token},
        )
    if response.status_code >= 400:
        raise HTTPException(response.status_code, response.text[:2000])
    return response.json()


@app.post("/v1/reviews", dependencies=[Depends(authenticate)])
async def create_review(payload: ReviewPayload, x_amas_api_token: Annotated[str, Header()]) -> Any:
    async with httpx.AsyncClient(timeout=60) as client:
        response = await client.post(
            f"{N8N_BASE_URL}{N8N_REVIEW_PATH}",
            headers={"X-AMAS-API-Token": x_amas_api_token},
            json=payload.model_dump(mode="json"),
        )
    if response.status_code >= 400:
        raise HTTPException(response.status_code, response.text[:2000])
    return response.json()


@app.post("/internal/reports", dependencies=[Depends(authenticate_internal)])
async def archive_generated_report(payload: GeneratedReportPayload) -> dict[str, Any]:
    encoded = json.dumps(payload.report, ensure_ascii=False, sort_keys=True, separators=(",", ":")).encode("utf-8")
    digest = hashlib.sha256(encoded).hexdigest()
    tenant = re.sub(r"[^A-Za-z0-9._-]+", "_", payload.tenant_id)[:80] or "default"
    key = f"reports/{tenant}/{payload.case_id}/v{payload.report_version:04d}-{digest}.json"
    s3_client().put_object(
        Bucket=S3_BUCKET,
        Key=key,
        Body=encoded,
        ContentType="application/json",
        Metadata={
            "case-id": str(payload.case_id),
            "report-id": str(payload.report_id),
            "report-version": str(payload.report_version),
            "sha256": digest,
        },
    )
    return {"storage_bucket": S3_BUCKET, "storage_key": key, "sha256": digest, "size_bytes": len(encoded)}
